import os
import win32com
from win32com.client import Dispatch
from docx import Document
import pandas as pd

def parse_docx(): # 使用python_docx模块处理docx文档
	global data
	doc = Document(PATH+'\\'+file)	# 打开一个docx“日报”文档
	for t in doc.tables:	# 遍历单个文档的所有表格
		try:
			for i in range(len(t.row_cells(0))):	# 遍历单个表格的第一行每一个数据格：
				if keyWord in t.row_cells(0)[i].text:	# 判断单个数据格值是否包含“手足口”关键字：如果有，就把该表格中的省市名称列和“手足口”数据列追加保存到数据集中
					fashengtime=[]
					diqu=[]
					fashengshu=[]
					fashengtime=[file] * (len(t.column_cells(0))-1)	# 使用文件名作为“报告发生时间”
					for j in range(1,len(t.column_cells(0))):
						diqu.append(t.cell(j,0).text.replace("\r","").replace("\x07",""))	# 读取“地区”信息
						fashengshu.append(t.cell(j,i).text.replace("\r","").replace("\x07",""))	# 读取“报告发生数”信息
					df_append=pd.DataFrame({"fashengtime":pd.Series(fashengtime),"diqu":pd.Series(diqu),"fashengshu":pd.Series(fashengshu)})	# 抽取到的数据
					data=data.append(df_append, ignore_index=True)	# 把抽取到的数据保存到数据集中
		except Exception as e:
			print(e)

def parse_doc():	# 使用win32com模块处理doc文档
	global data
	w = win32com.client.Dispatch('Word.Application')
	doc = w.Documents.Open( FileName = PATH+'\\'+file )	# 打开一个doc“日报”文档
	for t in doc.tables:	# 遍历单个文档的所有表格
		try:
			for i in range(len(t.Rows[0].Cells)):	# 遍历单个表格的第一行每一个数据格：
				if keyWord in t.Rows[0].Cells[i].Range.text:	# 判断单个数据格值是否包含“手足口”关键字：如果有，就把该表格中的省市名称列和“手足口”数据列追加保存到数据集中
					fashengtime=[]
					diqu=[]
					fashengshu=[]
					fashengtime=[file] * (len(t.Rows)-1)	# 使用文件名作为“报告发生时间”
					for j in range(1,len(t.Rows)):
						diqu.append(t.Rows[j].Cells[0].Range.text.replace("\r","").replace("\x07",""))	# 读取“地区”信息
						fashengshu.append(t.Rows[j].Cells[i].Range.text.replace("\r","").replace("\x07",""))	# 读取“报告发生数”信息
					df_append=pd.DataFrame({"fashengtime":pd.Series(fashengtime),"diqu":pd.Series(diqu),"fashengshu":pd.Series(fashengshu)})	# 从一个表格中抽取到的数据
					data=data.append(df_append, ignore_index=True)	# 把抽取到的数据保存到数据集中
		except Exception as e:
			print(e)
	doc.Close()	# 关闭打开的文档

# 以下为主代码
PATH = "F:\python\\word"	# 设置“日报”文档所在的文件夹路径
keyWord="手足口"	#设置查找的关键字
data=pd.DataFrame(columns=['fashengtime', 'diqu', 'fashengshu']) 	# 创建包含三列的数据集，三列分别为“报告发生时间”、“地区”和“报告发生数”，抽取出的数据追加到这里。
doc_files = os.listdir(PATH)	# 获取文件夹下所有“日报”文档的名称列表
for file in doc_files:	# 遍历“日报”文档名称列表
	if os.path.splitext(file)[1] == '.docx':	#如果是docx文档，用python_docx模块处理
		parse_docx()
	elif os.path.splitext(file)[1] == '.doc':	#如果是doc文档，用win32com模块处理
		parse_doc()
data.to_csv('F:\\python\\data.csv',encoding="utf_8_sig")	#把所有集成的数据保存为csv文件。